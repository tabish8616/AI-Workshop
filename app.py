import os
# MANDATORY FOR VERCEL: Headless backend and temp directory for matplotlib
os.environ["MPLCONFIGDIR"] = "/tmp"
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

import io
import json
import re
import zipfile
import base64
from datetime import datetime
from flask import Flask, render_template, request, jsonify, send_file
import pandas as pd

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, PageBreak
)

from google import genai
from google.genai import types

app = Flask(__name__)

# --- CONFIGURATION ---
# Note for Vercel Hobby tier: Serverless function limit is 10 seconds by default. 
# "maxDuration" in vercel.json handles limits on Vercel Pro.
MODEL_NAME = "gemini-2.5-flash-lite"

def call_gemini(prompt, json_mode=False, temperature=0.5, max_output_tokens=2048):
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        return None, "GEMINI_API_KEY environment variable is not configured."
    
    try:
        client = genai.Client(api_key=api_key)
        config = types.GenerateContentConfig(
            temperature=temperature,
            max_output_tokens=max_output_tokens,
        )
        if json_mode:
            config.response_mime_type = "application/json"
            
        response = client.models.generate_content(
            model=MODEL_NAME,
            contents=prompt,
            config=config
        )
        text = getattr(response, "text", None)
        if not text or not text.strip():
            return None, "The AI returned an empty response. Please try again."
        return text.strip(), None
    except Exception as e:
        return None, f"Gemini API error: {str(e)}"

def safe_json_parse(text):
    if not text:
        return None
    cleaned = text.strip()
    cleaned = re.sub(r"^```(json)?", "", cleaned, flags=re.IGNORECASE).strip()
    cleaned = re.sub(r"```$", "", cleaned).strip()
    try:
        return json.loads(cleaned)
    except Exception:
        match = re.search(r"(\[.*\]|\{.*\})", cleaned, re.DOTALL)
        if match:
            try:
                return json.loads(match.group(1))
            except Exception:
                return None
    return None

# --- EXPORT LOGIC ---
def journal_font(journal_format):
    mapping = {
        "IEEE": ("Times New Roman", 10),
        "Springer": ("Times New Roman", 11),
        "Elsevier": ("Times New Roman", 11),
        "College Project": ("Calibri", 12),
        "Generic Academic": ("Times New Roman", 12),
    }
    return mapping.get(journal_format, ("Times New Roman", 12))

def build_docx(paper, journal_format):
    font_name, font_size = journal_font(journal_format)
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(font_size)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(paper.get("Title", "Research Paper"))
    run.bold = True
    run.font.size = Pt(font_size + 6)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(f"{journal_format} Format · Generated {datetime.now().strftime('%B %Y')}")
    sub_run.italic = True
    sub_run.font.size = Pt(font_size - 1)
    doc.add_page_break()

    for key, content in paper.items():
        if key == "Title":
            continue
        heading = doc.add_heading(key, level=1)
        for r in heading.runs:
            r.font.color.rgb = RGBColor(0x0D, 0x3B, 0x66)
        for para in str(content).split("\n"):
            if para.strip():
                doc.add_paragraph(para.strip())
        doc.add_page_break()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def build_pdf(paper):
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=0.9 * inch, bottomMargin=0.9 * inch,
                             leftMargin=0.9 * inch, rightMargin=0.9 * inch)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("TitleStyle", parent=styles["Title"],
                                  textColor=colors.HexColor("#0D3B66"), alignment=TA_CENTER)
    heading_style = ParagraphStyle("HeadingStyle", parent=styles["Heading2"],
                                    textColor=colors.HexColor("#1565C0"), spaceBefore=14)
    body_style = ParagraphStyle("BodyStyle", parent=styles["BodyText"],
                                 alignment=TA_JUSTIFY, fontSize=10.5, leading=15)

    flow = [Paragraph(paper.get("Title", "Research Paper"), title_style), Spacer(1, 16)]
    for key, content in paper.items():
        if key == "Title":
            continue
        flow.append(Paragraph(key, heading_style))
        for para in str(content).split("\n"):
            if para.strip():
                safe = para.strip().replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                flow.append(Paragraph(safe, body_style))
                flow.append(Spacer(1, 6))
        flow.append(PageBreak())

    doc.build(flow)
    buf.seek(0)
    return buf

def bib_key(author, year):
    base = re.sub(r"[^A-Za-z]", "", str(author).split(",")[0].split(" ")[0]) or "Ref"
    return f"{base}{year}"

def build_latex_zip(paper, lit_table):
    title = paper.get("Title", "Research Paper").replace("&", "and")

    body_sections = []
    for key, content in paper.items():
        if key in ("Title", "References"):
            continue
        safe_content = str(content).replace("%", "\\%").replace("_", "\\_").replace("&", "\\&")
        body_sections.append(f"\\section{{{key}}}\n{safe_content}\n")

    cite_keys = [bib_key(r.get("author", "Author"), r.get("year", "n.d.")) for r in lit_table]
    cite_str = ", ".join(cite_keys) if cite_keys else ""

    main_tex = f"""\\documentclass[12pt]{{article}}
\\usepackage[margin=1in]{{geometry}}
\\usepackage{{cite}}
\\title{{{title}}}
\\author{{}}
\\date{{}}
\\begin{{document}}
\\maketitle

{"".join(body_sections)}

\\section{{References}}
This study draws on prior work \\cite{{{cite_str}}}. See references.bib.
All entries require manual verification.

\\bibliographystyle{{apalike}}
\\bibliography{{references}}
\\end{{document}}
"""

    bib_entries = []
    for row in lit_table:
        key = bib_key(row.get("author", "Author"), row.get("year", "n.d."))
        bib_entries.append(
            f"@article{{{key},\n"
            f"  author = {{{row.get('author','Unknown')}}},\n"
            f"  year = {{{row.get('year','n.d.')}}},\n"
            f"  title = {{{row.get('findings','Untitled')[:60]}}},\n"
            f"  note = {{Verification Required}}\n}}\n"
        )
    references_bib = "\n".join(bib_entries) if bib_entries else (
        "% No literature entries available.\n% Verification Required for all sources.\n"
    )

    readme = (
        "AI Research Studio - LaTeX Export\n"
        "==================================\n\n"
        "Files included:\n"
        "  main.tex        - Main paper source\n"
        "  references.bib  - Bibliography entries (AI-generated, unverified)\n\n"
        "To compile: run pdflatex main.tex followed by bibtex main, then pdflatex main.tex twice.\n\n"
        "IMPORTANT: All references and citations were generated by AI and must be manually\n"
        "verified against original academic sources before submission.\n"
    )

    zip_buf = io.BytesIO()
    with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("main.tex", main_tex)
        zf.writestr("references.bib", references_bib)
        zf.writestr("README.txt", readme)
    zip_buf.seek(0)
    return zip_buf

def build_questionnaire_docx(questions):
    doc = Document()
    doc.add_heading("Research Questionnaire", level=1)
    doc.add_paragraph("Please rate each statement on a 5-point scale: 1 = Strongly Disagree, 5 = Strongly Agree.")
    for i, q in enumerate(questions, start=1):
        doc.add_paragraph(f"{i}. {q}    [1] [2] [3] [4] [5]")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# --- ROUTES ---
@app.route("/")
def index():
    return render_template("index.html")

@app.route("/api/generate_research_plan", methods=["POST"])
def api_generate_research_plan():
    data = request.json
    prompt = f"""You are an academic research assistant.
Topic: {data.get('topic')}
Domain: {data.get('domain')}
Research Level: {data.get('level')}
Country: {data.get('country')}
Research Type: {data.get('rtype')}

Return ONLY valid JSON with exactly these keys:
{{"refined_topic": "...", "scope": "...", "keywords": ["...","..."], "suggested_direction": "..."}}
No markdown fences, no extra text."""
    text, err = call_gemini(prompt, json_mode=True, max_output_tokens=600)
    if err: return jsonify({"error": err}), 400
    parsed = safe_json_parse(text)
    if not parsed: return jsonify({"error": "Could not parse AI response."}), 400
    return jsonify(parsed)

@app.route("/api/generate_problem_statement", methods=["POST"])
def api_generate_problem_statement():
    data = request.json
    plan = data.get("plan", {})
    mode = data.get("mode", "generate")
    existing = data.get("existing", "")
    base = f"Refined Topic: {plan.get('refined_topic','')}\nScope: {plan.get('scope','')}\n"
    if mode == "generate": instr = "Write a concise academic problem statement (120-180 words)."
    elif mode == "regenerate": instr = "Write a fresh, differently-worded academic problem statement (120-180 words)."
    elif mode == "expand": instr = f"Expand this problem statement with more academic depth (200-260 words):\n{existing}"
    else: instr = f"Simplify this problem statement into clear beginner-friendly language (100-140 words):\n{existing}"
    
    prompt = base + instr + "\nReturn only the statement text, no heading."
    text, err = call_gemini(prompt, max_output_tokens=500)
    if err: return jsonify({"error": err}), 400
    return jsonify({"problem_statement": text})

@app.route("/api/generate_objectives", methods=["POST"])
def api_generate_objectives():
    data = request.json
    plan = data.get("plan", {})
    problem_statement = data.get("problem_statement", "")
    prompt = f"""Refined Topic: {plan.get('refined_topic','')}
Scope: {plan.get('scope','')}
Problem Statement: {problem_statement[:500]}

Generate academic research content in markdown with three headed sections:
### Research Objectives
(3-5 concise bullet points)
### Research Questions
(3-5 concise bullet points)
### Hypotheses
(2-4 bullet points, or state "Not applicable for qualitative research" if appropriate)
Return only markdown."""
    text, err = call_gemini(prompt, max_output_tokens=700)
    if err: return jsonify({"error": err}), 400
    return jsonify({"objectives_text": text})

@app.route("/api/generate_literature_table", methods=["POST"])
def api_generate_literature_table():
    data = request.json
    plan = data.get("plan", {})
    prompt = f"""Refined Topic: {plan.get('refined_topic','')}
Domain Scope: {plan.get('scope','')}

Generate 7 illustrative academic literature review entries relevant to this topic.
Return ONLY a JSON array, each item with exactly these keys:
"author", "year", "methodology", "findings", "limitation".
Keep each field under 25 words. No markdown fences."""
    text, err = call_gemini(prompt, json_mode=True, max_output_tokens=1200)
    if err: return jsonify({"error": err}), 400
    parsed = safe_json_parse(text)
    if not isinstance(parsed, list): return jsonify({"error": "Failed to parse literature JSON."}), 400
    return jsonify({"lit_table": parsed})

@app.route("/api/generate_literature_narrative", methods=["POST"])
def api_generate_literature_narrative():
    data = request.json
    table = data.get("lit_table", [])
    compact = "\n".join(f"- {row.get('author','')} ({row.get('year','')}): {row.get('findings','')}" for row in table)
    prompt = f"""Studies:
{compact}

Write a cohesive academic narrative literature review (250-350 words) synthesizing
these studies, noting patterns, agreements and contrasts. Return only prose, no headings."""
    text, err = call_gemini(prompt, max_output_tokens=900)
    if err: return jsonify({"error": err}), 400
    return jsonify({"lit_narrative": text})

@app.route("/api/generate_research_gap", methods=["POST"])
def api_generate_research_gap():
    data = request.json
    narrative = data.get("lit_narrative", "")[:900]
    prompt = f"""Literature Review Summary:
{narrative}

Identify research gaps. Return markdown with these headings:
### Academic Gap
### Practical Gap
### Future Research Opportunity
2-3 sentences under each heading."""
    text, err = call_gemini(prompt, max_output_tokens=600)
    if err: return jsonify({"error": err}), 400
    return jsonify({"research_gap": text})

@app.route("/api/generate_methodology", methods=["POST"])
def api_generate_methodology():
    data = request.json
    plan = data.get("plan", {})
    m = data.get("methodology_inputs", {})
    prompt = f"""Refined Topic: {plan.get('refined_topic','')}
Approach: {m.get('approach')}
Design: {m.get('design')}
Sample Size: {m.get('sample_size')}
Sampling Method: {m.get('sampling_method')}
Data Collection Method: {m.get('data_collection')}

Write a complete academic methodology section (300-400 words) covering research design,
population and sample, sampling technique, data collection procedure, and data analysis
approach. Use markdown subheadings. Return only markdown."""
    text, err = call_gemini(prompt, max_output_tokens=900)
    if err: return jsonify({"error": err}), 400
    return jsonify({"methodology_text": text})

@app.route("/api/generate_questionnaire", methods=["POST"])
def api_generate_questionnaire():
    data = request.json
    plan = data.get("plan", {})
    m = data.get("methodology_inputs", {})
    prompt = f"""Topic: {plan.get('refined_topic','')}
Study Approach: {m.get('approach')}

Generate 12 five-point Likert scale questionnaire items (Strongly Disagree to Strongly Agree)
measuring constructs relevant to this topic. Return ONLY a JSON array of 12 short statement
strings (no numbering, statement form, no question marks). No markdown fences."""
    text, err = call_gemini(prompt, json_mode=True, max_output_tokens=700)
    if err: return jsonify({"error": err}), 400
    parsed = safe_json_parse(text)
    if not isinstance(parsed, list): return jsonify({"error": "Failed to parse questionnaire JSON."}), 400
    return jsonify({"questionnaire": parsed})

@app.route("/api/generate_results", methods=["POST"])
def api_generate_results():
    data = request.json
    plan = data.get("plan", {})
    methodology_text = data.get("methodology_text", "")
    prompt = f"""Refined Topic: {plan.get('refined_topic','')}
Methodology Summary: {methodology_text[:400]}

Generate markdown with these headings, 60-100 words each:
### Expected Findings
### Managerial Implications
### Academic Implications
Return only markdown."""
    text, err = call_gemini(prompt, max_output_tokens=700)
    if err: return jsonify({"error": err}), 400
    return jsonify({"results_text": text})

@app.route("/api/generate_conclusion", methods=["POST"])
def api_generate_conclusion():
    data = request.json
    plan = data.get("plan", {})
    results_text = data.get("results_text", "")
    prompt = f"""Refined Topic: {plan.get('refined_topic','')}
Expected Results Summary: {results_text[:500]}

Generate markdown with these headings, 60-100 words each:
### Conclusion
### Recommendations
### Limitations
### Future Scope
Return only markdown."""
    text, err = call_gemini(prompt, max_output_tokens=800)
    if err: return jsonify({"error": err}), 400
    return jsonify({"conclusion_text": text})

@app.route("/api/generate_references", methods=["POST"])
def api_generate_references():
    data = request.json
    plan = data.get("plan", {})
    style = data.get("citation_style", "APA")
    table = data.get("lit_table", [])
    authors = "; ".join(f"{r.get('author','')} ({r.get('year','')})" for r in table) or "general sources"
    prompt = f"""Topic: {plan.get('refined_topic','')}
Citation Style: {style}
Known authors referenced earlier: {authors}

Generate a reference list of 10 entries in {style} citation style relevant to this topic.
Return as a markdown numbered list, one reference per line, no extra commentary."""
    text, err = call_gemini(prompt, max_output_tokens=900)
    if err: return jsonify({"error": err}), 400
    lines = []
    for line in text.split("\n"):
        line = line.strip()
        if not line: continue
        if "[Verification Required]" not in line:
            line = line + " [Verification Required]"
        lines.append(line)
    return jsonify({"references_text": "\n".join(lines)})

@app.route("/api/progress_chart", methods=["POST"])
def api_progress_chart():
    data = request.json
    completed = data.get("completed", 0)
    total = data.get("total", 10)
    
    fig, ax = plt.subplots(figsize=(2.4, 2.4))
    remaining = total - completed
    colors_ = ["#1565C0", "#DCE8F7"]
    ax.pie([completed, remaining] if remaining > 0 else [completed, 0.0001],
           colors=colors_, startangle=90,
           wedgeprops=dict(width=0.35, edgecolor="white"))
    ax.text(0, 0, f"{completed}/{total}", ha="center", va="center",
            fontsize=15, fontweight="bold", color="#0D3B66")
    ax.set_aspect("equal")
    fig.patch.set_alpha(0)
    
    buf = io.BytesIO()
    plt.savefig(buf, format="png", bbox_inches='tight', transparent=True)
    plt.close(fig)
    img_b64 = base64.b64encode(buf.getvalue()).decode("utf-8")
    return jsonify({"image": f"data:image/png;base64,{img_b64}"})

# --- DOWNLOAD ENDPOINTS ---
@app.route("/api/download/docx", methods=["POST"])
def download_docx():
    data = request.json
    buf = build_docx(data.get("paper", {}), data.get("journal_format", "Generic Academic"))
    return send_file(buf, as_attachment=True, download_name="research_paper.docx", mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

@app.route("/api/download/pdf", methods=["POST"])
def download_pdf():
    data = request.json
    buf = build_pdf(data.get("paper", {}))
    return send_file(buf, as_attachment=True, download_name="research_paper.pdf", mimetype="application/pdf")

@app.route("/api/download/latex", methods=["POST"])
def download_latex():
    data = request.json
    buf = build_latex_zip(data.get("paper", {}), data.get("lit_table", []))
    return send_file(buf, as_attachment=True, download_name="research_paper_latex.zip", mimetype="application/zip")

@app.route("/api/download/questionnaire", methods=["POST"])
def download_questionnaire():
    data = request.json
    buf = build_questionnaire_docx(data.get("questionnaire", []))
    return send_file(buf, as_attachment=True, download_name="questionnaire.docx", mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

if __name__ == "__main__":
    app.run(debug=True, port=5000)